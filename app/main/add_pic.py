from docx import Document
from docx.shared import Cm

def addPic(doc_name,pic_name):
	doc_name='docs/'+doc_name
	document=Document(doc_name)
	tables=document.tables
	print(len(tables))
	table=tables[0]
	cells=table._cells
	print(len(cells))
	cells_string=[cell.text for cell in cells]
	# print(cells_string)
	col_num= len(table.columns)
	print(col_num)
	row_num=len(table.rows)
	print(row_num)

	row0=table.rows[0]
	col0=table.columns[7]

	row0_str=[cell.text for cell in row0.cells]
	print(row0_str)
	col0_str=[cell.text for cell in col0.cells]
	print(col0_str)
	print(table.cell(0,7).text)
	pic=table.cell(0,6)
	run=pic.paragraphs[0].add_run()
	run.add_picture('docs/'+pic_name,width=Cm(2.5))
	print('docs/'+pic_name+"   "+doc_name)

	document.save(doc_name)
